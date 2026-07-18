from __future__ import annotations

import hashlib
import re
from collections import OrderedDict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Youss\OneDrive\Documents\DGOP")
SOURCE = ROOT / "_analysis" / "v4_sprint_update_text.txt"
WORK_DIR = ROOT / "_analysis" / "sprint_word_template_work"
REFERENCE = Path(
    r"C:\Users\Youss\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.0\skills\artifact-template-system-design\assets\reference.docx"
)
FINAL_DOCX = ROOT / "DGOP_All_Sprints_System_Design_Roadmap.docx"
ARTIFACT_MD = WORK_DIR / "artifact.md"

BLUE = "082A4A"
MID_BLUE = "5B7085"
LIGHT_BLUE = "D8EAF8"
PALE_BLUE = "F3F8FC"
GRID = "D9E4ED"
TEXT = "1F2A33"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def clean(text: str) -> str:
    text = text.replace("\u2013", "-").replace("\u2014", "-").replace("\u00a0", " ")
    return re.sub(r"\s+", " ", text).strip()


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def format_table(table, header=True) -> None:
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = True
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            mark_header_row(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if header and r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(BLUE)
            else:
                set_cell_shading(cell, PALE_BLUE if r_idx % 2 == 0 else "FFFFFF")


def add_para(doc: Document, text: str = "", style: str = "normal", bold: bool = False, color: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_lead(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph(style="normal")
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p.add_run(text)


def _max_numbering_id(numbering, tag_name: str, attr_name: str) -> int:
    values = []
    for node in numbering.findall(qn(tag_name)):
        value = node.get(qn(attr_name))
        if value and value.isdigit():
            values.append(int(value))
    return max(values) if values else 0


def ensure_bullet_numbering(doc: Document) -> int:
    existing = getattr(doc, "_dgop_bullet_num_id", None)
    if existing is not None:
        return existing
    numbering = doc.part.numbering_part.element
    abstract_id = _max_numbering_id(numbering, "w:abstractNum", "w:abstractNumId") + 1
    num_id = _max_numbering_id(numbering, "w:num", "w:numId") + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    ppr.append(ind)
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    rpr.append(fonts)
    for child in (start, fmt, text, jc, ppr, rpr):
        lvl.append(child)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    setattr(doc, "_dgop_bullet_num_id", num_id)
    return num_id


def apply_bullet(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_bullets(doc: Document, items: list[str]) -> None:
    num_id = ensure_bullet_numbering(doc)
    for item in items:
        if not item:
            continue
        p = doc.add_paragraph(style="normal")
        apply_bullet(p, num_id)
        p.paragraph_format.space_after = Pt(1.5)
        p.add_run(item)


def extract_sections(block: str) -> OrderedDict[str, list[str]]:
    sections: OrderedDict[str, list[str]] = OrderedDict()
    current = "Overview"
    sections[current] = []
    for raw in block.splitlines():
        line = clean(raw)
        if not line:
            continue
        is_label = line.endswith(":") and 2 <= len(line[:-1]) <= 70 and not line.startswith("http")
        if is_label:
            current = line[:-1]
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)
    return sections


def parse_source() -> tuple[list[dict], OrderedDict[str, list[str]]]:
    text = SOURCE.read_text(encoding="utf-8")
    cutoff = text.find("Definition Of Done For Every Sprint")
    sprint_region = text[:cutoff] if cutoff != -1 else text
    tail = text[cutoff:] if cutoff != -1 else ""

    release_matches = list(re.finditer(r"(?m)^Release\s+(\d+):\s*(.+)$", sprint_region))
    sprint_matches = list(re.finditer(r"(?m)^Sprint\s+(\d+):\s*(.+)$", sprint_region))

    def release_for(pos: int) -> str:
        current = "Roadmap"
        for rel in release_matches:
            if rel.start() <= pos:
                current = f"Release {rel.group(1)}: {rel.group(2)}"
            else:
                break
        return current

    sprints: list[dict] = []
    for idx, match in enumerate(sprint_matches):
        start = match.end()
        end = sprint_matches[idx + 1].start() if idx + 1 < len(sprint_matches) else len(sprint_region)
        block = sprint_region[start:end].strip()
        sections = extract_sections(block)
        goal = ""
        if sections.get("Overview") and sections["Overview"][0].lower().startswith("goal:"):
            goal = clean(sections["Overview"].pop(0)[5:])
        elif sections.get("Goal"):
            goal = " ".join(sections["Goal"])
        sprints.append(
            {
                "number": int(match.group(1)),
                "title": clean(match.group(2)),
                "release": release_for(match.start()),
                "goal": goal,
                "sections": sections,
            }
        )

    appendix_headings = [
        "Definition Of Done For Every Sprint",
        "Cross-Cutting Engineering Standards",
        "Suggested First Backlog Epics",
        "Immediate Implementation Sequence After Plan Approval",
        "Key Risks And Mitigations",
    ]
    appendices: OrderedDict[str, list[str]] = OrderedDict()
    for i, heading in enumerate(appendix_headings):
        start = tail.find(heading)
        if start == -1:
            continue
        next_positions = [tail.find(h, start + len(heading)) for h in appendix_headings[i + 1 :]]
        next_positions = [p for p in next_positions if p != -1]
        end = min(next_positions) if next_positions else len(tail)
        body = tail[start + len(heading) : end]
        appendices[heading] = [clean(line) for line in body.splitlines() if clean(line)]
    return sprints, appendices


def add_cover(doc: Document, sprints: list[dict]) -> None:
    for _ in range(6):
        add_para(doc, "")
    title1 = doc.add_paragraph(style="Title")
    title1.add_run("DGOP").font.color.rgb = RGBColor.from_string(MID_BLUE)
    title2 = doc.add_paragraph(style="Title")
    r = title2.add_run("All Sprints System Design Roadmap")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    add_para(doc, "Enterprise Governance Operations Platform sprint catalog and delivery architecture.", "normal")
    add_para(doc, "")

    meta = doc.add_table(rows=1, cols=3)
    values = [
        ("STATUS", "Draft / Working Roadmap"),
        ("OWNER", "DGOP Team"),
        ("LAST UPDATED", "July 16, 2026"),
    ]
    for cell, (label, value) in zip(meta.rows[0].cells, values):
        p = cell.paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(7)
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(f"\n{value}")
    format_table(meta, header=False)
    add_para(doc, "")
    info = doc.add_table(rows=4, cols=2)
    rows = [
        ("Authors", "Codex with DGOP source sprint plan"),
        ("Reviewers", "DGOP senior stakeholders"),
        ("Related docs", "_analysis/v4_sprint_update_text.txt; DGOP v4/v5 planning material"),
        ("Scope", f"Sprints {sprints[0]['number']} through {sprints[-1]['number']} with release-level delivery standards."),
    ]
    for row, (k, v) in zip(info.rows, rows):
        row.cells[0].text = k
        row.cells[1].text = v
    format_table(info, header=False)
    doc.add_page_break()


def add_summary_sections(doc: Document, sprints: list[dict]) -> None:
    doc.add_heading("1. Abstract", level=1)
    add_para(
        doc,
        "This document consolidates the DGOP sprint roadmap into a system-design style delivery artifact. It covers the governance platform foundation, operational modules, transparency and privacy workflows, integration hardening, production readiness, intelligence improvements, and go-live handover.",
    )
    add_para(
        doc,
        "The roadmap is organized by release waves so senior stakeholders can understand what is delivered, why it matters, and how the work connects to backend services, Angular workspaces, compliance controls, and user-facing operating flows.",
    )

    doc.add_heading("2. Goals and Non-Goals", level=1)
    table = doc.add_table(rows=5, cols=2)
    rows = [
        ("Goals", "Non-goals"),
        ("Keep all planned DGOP sprints in one readable Word artifact.", "This document does not replace detailed QA test cases or source code tickets."),
        ("Show each sprint goal, work areas, acceptance criteria, and release wave.", "This document does not claim every planned future sprint is fully implemented."),
        ("Use a system-design format that connects roadmap work to architecture and operations.", "This document does not expose secrets, credentials, or local environment values."),
        ("Preserve the shared DGOP definition of done and cross-cutting standards.", "This document does not change the current product scope without senior approval."),
    ]
    for row, values in zip(table.rows, rows):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]
    format_table(table)

    doc.add_heading("3. Roadmap Architecture", level=1)
    releases: OrderedDict[str, list[dict]] = OrderedDict()
    for sprint in sprints:
        releases.setdefault(sprint["release"], []).append(sprint)
    rel_table = doc.add_table(rows=1, cols=4)
    headers = ["Release", "Sprint Range", "Theme", "Outcome"]
    for cell, text in zip(rel_table.rows[0].cells, headers):
        cell.text = text
    for release, sprint_rows in releases.items():
        row = rel_table.add_row().cells
        nums = [s["number"] for s in sprint_rows]
        row[0].text = release.split(":", 1)[0]
        row[1].text = f"{min(nums)}-{max(nums)}"
        row[2].text = release.split(":", 1)[1].strip() if ":" in release else release
        row[3].text = sprint_rows[-1]["goal"] or sprint_rows[-1]["title"]
    format_table(rel_table)

    doc.add_heading("4. Sprint Delivery Lifecycle", level=1)
    steps = [
        "Scope the sprint against DGOP operating towers, personas, NDI/PDP/security obligations, and demo goals.",
        "Implement backend APIs, database rules, authorization, audit logging, and workflow hooks before UI-only polish.",
        "Build Angular screens with shared design-system components, bilingual copy, light/dark parity, and RTL support.",
        "Verify with focused unit/API tests, UI checks, accessibility review, and sprint-level QA evidence.",
        "Document delivered, deferred, and mocked capabilities before moving to the next sprint.",
    ]
    add_bullets(doc, steps)

    doc.add_heading("5. Sprint Catalog Overview", level=1)
    catalog = doc.add_table(rows=1, cols=4)
    for cell, text in zip(catalog.rows[0].cells, ["Sprint", "Release", "Title", "Goal"]):
        cell.text = text
    for sprint in sprints:
        row = catalog.add_row().cells
        row[0].text = str(sprint["number"])
        row[1].text = sprint["release"].split(":", 1)[0]
        row[2].text = sprint["title"]
        row[3].text = sprint["goal"] or "See detailed section."
    format_table(catalog)


def add_sprint_details(doc: Document, sprints: list[dict]) -> None:
    doc.add_page_break()
    doc.add_heading("6. Detailed Sprint Catalog", level=1)
    current_release = None
    for sprint in sprints:
        if sprint["release"] != current_release:
            current_release = sprint["release"]
            doc.add_heading(current_release, level=2)
        doc.add_heading(f"Sprint {sprint['number']}: {sprint['title']}", level=3)
        if sprint["goal"]:
            add_lead(doc, "Goal", sprint["goal"])
        for label, items in sprint["sections"].items():
            if label == "Overview" and not items:
                continue
            if label == "Goal":
                continue
            heading = label if label != "Overview" else "Additional notes"
            p = doc.add_paragraph(style="normal")
            r = p.add_run(heading)
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(MID_BLUE)
            add_bullets(doc, items)


def add_appendices(doc: Document, appendices: OrderedDict[str, list[str]]) -> None:
    doc.add_page_break()
    doc.add_heading("7. Roadmap-Wide Standards and Appendices", level=1)
    for heading, items in appendices.items():
        doc.add_heading(heading, level=2)
        add_bullets(doc, items)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    for name in ("Normal", "normal"):
        if name in doc.styles:
            style = doc.styles[name]
            style.font.name = "Helvetica Neue"
            style.font.size = Pt(9.2)
            style.font.color.rgb = RGBColor.from_string(TEXT)
            style.paragraph_format.space_after = Pt(4)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        if name in doc.styles:
            style = doc.styles[name]
            style.font.name = "Helvetica Neue"
            style.font.color.rgb = RGBColor.from_string(BLUE if name != "Heading 3" else MID_BLUE)
            style.font.bold = True
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(4)


def write_artifact(sprints: list[dict], appendices: OrderedDict[str, list[str]]) -> None:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    ARTIFACT_MD.write_text(
        "\n".join(
            [
                "# DGOP Sprint Roadmap System Design Artifact",
                "",
                f"Reference: {REFERENCE}",
                f"Reference SHA-256: {sha256(REFERENCE)}",
                "Reference page system: US Letter portrait; 0.70in left/right/top, 0.62in bottom; different first page.",
                "Reference style system: Helvetica Neue family, dark blue title/heading hierarchy, pale-blue metadata and comparison tables.",
                f"Content source: {SOURCE}",
                f"Generated output: {FINAL_DOCX}",
                f"Sprint coverage: {sprints[0]['number']} through {sprints[-1]['number']} ({len(sprints)} sprints).",
                f"Appendix sections: {', '.join(appendices.keys())}",
                "Editable slots: cover title/subtitle/metadata; abstract; goals/non-goals table; roadmap architecture table; delivery lifecycle bullets; detailed sprint catalog; appendix standards.",
                "Fidelity note: body content is replaced intentionally while retaining the retained reference document's section setup, styles, and blue table treatment.",
                "Render note: PNG render QA requires LibreOffice/soffice; if unavailable, structural DOCX audits are used.",
                "",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    sprints, appendices = parse_source()
    write_artifact(sprints, appendices)
    doc = Document(REFERENCE)
    clear_body(doc)
    style_document(doc)
    add_cover(doc, sprints)
    add_summary_sections(doc, sprints)
    add_sprint_details(doc, sprints)
    add_appendices(doc, appendices)
    doc.core_properties.title = "DGOP All Sprints System Design Roadmap"
    doc.core_properties.subject = "DGOP sprint roadmap and delivery architecture"
    doc.core_properties.author = "Codex"
    doc.core_properties.last_modified_by = "Codex"
    doc.core_properties.keywords = "DGOP, sprint roadmap, system design, data governance"
    doc.core_properties.comments = "Created from the retained System Design template reference."
    doc.save(FINAL_DOCX)
    print(FINAL_DOCX)
    print(f"sprints={len(sprints)} appendices={len(appendices)}")


if __name__ == "__main__":
    main()
