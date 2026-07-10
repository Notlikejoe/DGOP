from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"C:\Users\Youss\OneDrive\Documents\DGOP\DGOP Functionality Changes Since UI Redesign.docx"


def set_run_font(run, name="Calibri", size=11, color=None, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(paragraph, before=16, after=8, line=1.1)
        size, color = 16, "2E74B5"
    elif level == 2:
        set_paragraph_spacing(paragraph, before=12, after=6, line=1.1)
        size, color = 13, "2E74B5"
    else:
        set_paragraph_spacing(paragraph, before=8, after=4, line=1.1)
        size, color = 12, "1F4D78"
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=6, line=1.1)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, before=0, after=4, line=1.167)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    tc_pr.append(shd)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.1)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color="1F3A5F", bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("DGOP functionality change log")
    set_run_font(run, size=9, color="666666")


doc = Document()
configure_document(doc)

title = doc.add_paragraph()
set_paragraph_spacing(title, before=0, after=3, line=1.1)
run = title.add_run("DGOP Functionality Changes Since UI Redesign")
set_run_font(run, size=22, color="0B2545", bold=True)

subtitle = doc.add_paragraph()
set_paragraph_spacing(subtitle, before=0, after=12, line=1.1)
run = subtitle.add_run("Functional change log through July 6, 2026")
set_run_font(run, size=11, color="555555")

add_callout(
    doc,
    "Scope",
    "This document lists functional capabilities added or hardened after the UI redesign phase and during the V4 sprint update work. Visual-only changes such as spacing, colors, card styling, and layout polish are intentionally excluded.",
)

add_heading(doc, "Planning And Sprint Alignment", 1)
for item in [
    "Created the V4 sprint update document from the enhanced V4 specification package.",
    "Aligned the product direction around the updated 12-sprint roadmap.",
    "Prioritized Sprint 12 awareness operations and Sprint 13 data quality operations for implementation.",
]:
    add_bullet(doc, item)

add_heading(doc, "Foundation And Navigation Capabilities", 1)
for item in [
    "Added an authenticated About DGOP section that explains the platform purpose and operating model.",
    "Added bilingual platform copy for new governance, training, and data-quality areas.",
    "Added new governed routes and navigation entries for Training and Awareness and Data Quality Operations.",
    "Kept new sections permission-gated so users only see capabilities allowed by their roles.",
]:
    add_bullet(doc, item)

add_heading(doc, "Dashboard And Executive Signals", 1)
for item in [
    "Added training readiness signals to the dashboard summary.",
    "Added data quality issue signals to the dashboard summary.",
    "Expanded risk totals to include overdue or expired training and critical or overdue data quality issues.",
    "Updated dashboard training calculations so expired completions no longer inflate readiness.",
    "Updated dashboard data-quality counts so they follow the same data-scope rules as the Data Quality workspace.",
]:
    add_bullet(doc, item)

add_heading(doc, "Sprint 12: Training And Awareness Operations", 1)
for item in [
    "Added a training course catalogue with tier, category, duration, delivery method, validity, and prerequisites.",
    "Added prerequisite validation so courses cannot depend on missing courses, themselves, or circular chains.",
    "Added role-based training requirements with configurable mandatory status and due days.",
    "Added training assignment lifecycle support for assigned, in progress, completed, expired, and waived records.",
    "Added role-assignment synchronization that creates missing training assignments for users who hold required roles.",
    "Added effective expiry handling so completed assignments past expiry read as expired without mutating read endpoints.",
    "Added certification tracks and certification attempts with current, renewal-due, expired, and non-passed states.",
    "Added continuing education activity tracking and CE hour rollups.",
    "Added community knowledge articles for steward awareness and reusable guidance.",
    "Added expert profiles and mentorship pair tracking for steward development.",
    "Added optional pagination for high-growth awareness lists such as assignments, certification attempts, CE activities, community articles, experts, and mentorships.",
]:
    add_bullet(doc, item)

add_heading(doc, "Sprint 13: Data Quality Operations", 1)
for item in [
    "Added a data quality issue register with status, severity, dimension, due date, linked asset, responsible person, and resolution summary.",
    "Added scope-aware summary metrics for total, open, critical, overdue, closed, and closure rate.",
    "Added issue list, detail, create, update, close, delete, and CSV import operations.",
    "Added automatic responsible-steward resolution from approved direct assignments and assignment rules.",
    "Added workflow case and remediation task creation when a data quality issue is registered.",
    "Added issue evidence/history trail entries for creation, updates, workflow creation, and closure.",
    "Added close behavior that completes open remediation tasks and closes the linked workflow case.",
    "Added transactional writes so issue, evidence, workflow, and audit records stay consistent together.",
    "Added optional pagination for the issue list so imports or growing issue history do not force unbounded responses.",
]:
    add_bullet(doc, item)

add_heading(doc, "Backend Stability And Governance Controls", 1)
for item in [
    "Added Prisma schema models, migrations, and seed data for training, awareness, certifications, mentorship, and data quality.",
    "Registered the Training and Data Quality modules in the API application module.",
    "Added DTO date validation for training, certification, CE, mentorship, and data quality due-date fields.",
    "Added cross-platform database and development scripts so local setup works without Bash-only commands.",
    "Added direct Data Quality service tests and included them in the normal API test command.",
    "Expanded dashboard tests for the new training and data-quality sections.",
    "Verified API health, database connectivity, authenticated dashboard summary, Training endpoints, and Data Quality endpoints after the fixes.",
]:
    add_bullet(doc, item)

add_heading(doc, "Current Verification Status", 1)
for item in [
    "API unit tests pass, including the new Data Quality and Training tests.",
    "API build passes.",
    "Prisma schema validation passes.",
    "Local migration status reports the database schema is up to date.",
    "Runtime health check reports the API as ok and the database as up.",
]:
    add_bullet(doc, item)

doc.save(OUT)
print(OUT)
