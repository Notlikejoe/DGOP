from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\Youss\Downloads\DGOP Sprints Details.docx")
OUTPUT = Path(r"C:\Users\Youss\OneDrive\Documents\DGOP\DGOP Sprint Details v4.0 Update.docx")


def compact(text: str) -> str:
    return " ".join((text or "").split())


def set_run_font(run, *, bold=None, size=None, name="Times New Roman"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = size


def replace_text(paragraph: Paragraph, text: str):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def paragraph_after(paragraph: Paragraph, text: str, template: Paragraph | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if template is not None:
        new_para.style = template.style
        new_para.paragraph_format.space_before = template.paragraph_format.space_before
        new_para.paragraph_format.space_after = template.paragraph_format.space_after
        new_para.paragraph_format.left_indent = template.paragraph_format.left_indent
        new_para.paragraph_format.right_indent = template.paragraph_format.right_indent
        new_para.paragraph_format.first_line_indent = template.paragraph_format.first_line_indent
        new_para.paragraph_format.line_spacing = template.paragraph_format.line_spacing
    run = new_para.add_run(text)
    set_run_font(run)
    return new_para


def find_first(doc: Document, text: str, *, start=0, end=None) -> int:
    end = len(doc.paragraphs) if end is None else end
    needle = compact(text)
    for i in range(start, end):
        if compact(doc.paragraphs[i].text) == needle:
            return i
    raise ValueError(f"Could not find paragraph: {text!r}")


def find_sprint_bounds(doc: Document, sprint_title: str) -> tuple[int, int]:
    start = find_first(doc, sprint_title)
    end = len(doc.paragraphs)
    for i in range(start + 1, len(doc.paragraphs)):
        text = compact(doc.paragraphs[i].text)
        if re.match(r"^Sprint \d+:", text) or re.match(r"^Release \d+:", text) or text == "Definition Of Done For Every Sprint":
            end = i
            break
    return start, end


def insert_after_exact(doc: Document, marker: str, lines: list[str], *, template_offset=1):
    idx = find_first(doc, marker)
    template_idx = min(idx + template_offset, len(doc.paragraphs) - 1)
    template = doc.paragraphs[template_idx]
    anchor = doc.paragraphs[idx]
    for line in reversed(lines):
        paragraph_after(anchor, line, template)


def insert_after_within_sprint(doc: Document, sprint_title: str, marker: str, lines: list[str]):
    start, end = find_sprint_bounds(doc, sprint_title)
    idx = find_first(doc, marker, start=start, end=end)
    template = doc.paragraphs[min(idx + 1, end - 1)]
    anchor = doc.paragraphs[idx]
    for line in reversed(lines):
        paragraph_after(anchor, line, template)


def replace_line(doc: Document, old: str, new: str):
    idx = find_first(doc, old)
    replace_text(doc.paragraphs[idx], new)


def main():
    doc = Document(SOURCE)

    replace_line(doc, "DGOP", "DGOP v4.0 Sprint Details Update")
    replace_line(
        doc,
        "Delivery length: 36 sprints, aligned to the original 18-month roadmap.",
        "Delivery length: 36 sprints, preserving the original 18-month cadence while absorbing the v4.0 enhanced scope.",
    )
    replace_line(
        doc,
        "Architecture: Angular frontend, Node.js backend, local PostgreSQL for development.",
        "Architecture: Angular frontend, NestJS/Node.js backend, PostgreSQL for development, and v4.0-aligned service boundaries for workflow, evidence, training, DQ, profiling, masking, ABAC, reporting, certification, and community capabilities.",
    )
    replace_line(
        doc,
        "First implementation style: modular monolith backend. Keep clear module boundaries so services can be split later if needed.",
        "First implementation style: modular monolith backend for speed and safety, with explicit bounded contexts so high-load or regulated services can be split later without redesign.",
    )
    replace_line(
        doc,
        "Local file storage can be used first for evidence uploads; object storage such as MinIO can be added later.",
        "Local file storage can be used first for evidence uploads; object storage such as MinIO/S3-compatible storage is planned for evidence packs, audit exports, and secure attachment retention.",
    )
    replace_line(
        doc,
        "External integrations start as mock adapters, CSV import/export, and documented API contracts before connecting to real tools.",
        "External integrations start as mock adapters, CSV import/export, and documented API contracts before connecting to catalog, DQ, DLP, ITSM, LMS, SIEM, Open Data portal, FOI channels, profiling, masking, and identity tools.",
    )

    insert_after_exact(
        doc,
        "Planning Assumptions",
        [
            "Source baseline: DGOP v4.0 Volume 1 Executive & Business Design, Volume 2 Functional Design, Volume 3 NDI Domain Models, Volume 4 Technical Design, Volume 5 Security & Compliance, and Volume 6 Procurement & Delivery.",
            "Scope update: v4.0 adds stewardship-role enhancement, charter and policy lifecycle management, decision-right timeframes, DGPO sizing, data-domain councils, certification, community of practice, expanded maturity, business result metrics, continuous improvement, DQ profiling, severity SLAs, escalation, RCA, rule lifecycle, MDM matching, business glossary, business lineage, data masking, recurring compliance, and role-data mapping.",
            "Delivery approach: enterprise-grade v4.0 capabilities are introduced in controlled increments so early sprints remain demonstrable and later sprints harden the platform for regulated production use.",
        ],
    )

    insert_after_exact(
        doc,
        "Core Backend Modules",
        [
            "charters_and_policies: governance charters, policy hierarchy, approval lifecycle, versioning, publication, and monitoring.",
            "decision_rights: RACI, decision matrix, escalation timeframes, authority levels, and approval accountability.",
            "certification: CDS, SDS, MDS tracks, exams, renewal, continuing education hours, and certification privileges.",
            "community: forums, knowledge base, best-practice library, expert directory, contribution incentives, and mentorship tracking.",
            "quality: DQ dimensions, scorecards, rules, profiling results, issues, RCA, remediation, validation, and closure evidence.",
            "mdm_glossary_lineage: master data matching, survivorship, business data elements, glossary approvals, and business lineage maps.",
            "security_policy: role-data mappings, ABAC policies, masking rules, Privacy by Design gates, and access certifications.",
            "compliance_calendar: recurring compliance tasks, scheduled reviews, reminders, escalations, and overdue penalties.",
        ],
    )

    insert_after_exact(
        doc,
        "Core Frontend Areas",
        [
            "Governance operating model workspace: charters, policy lifecycle, decision rights, RACI, councils, and stewardship operating views.",
            "Awareness operations workspace: training catalog, certification progress, continuing education, community, knowledge base, and mentorship.",
            "DQ operations workspace: 6-dimension scorecards, profiling summaries, rule lifecycle, issue remediation, RCA templates, and SLA escalation.",
            "Security/compliance workspace: role-data mapping designer, data masking policies, ABAC previews, access certification, recurring compliance calendar, and Privacy by Design gate reviews.",
            "Business value workspace: maturity heat maps, improvement plans, user survey results, ROI attribution, data asset valuation, glossary, and business lineage.",
        ],
    )

    insert_after_exact(
        doc,
        "Data Build Sequence",
        [
            "Charters, policy hierarchy, decision rights, RACI matrices, DGPO sizing, and domain councils.",
            "Training tiers, certification tracks, continuing education records, community content, experts, and mentorship pairs.",
            "DQ dimensions, DQ rules, profiling summaries, issue remediation, RCA evidence, severity SLAs, and escalation records.",
            "Role-data mappings, ABAC policy decisions, masking policies, Privacy by Design gates, and recurring compliance tasks.",
            "Business glossary terms, business lineage maps, MDM matching candidates, survivorship rules, and golden records.",
            "Maturity assessments, gap heat maps, improvement plans, business result metrics, surveys, and ROI attribution records.",
        ],
    )

    additions = {
        "Sprint 0: Project Setup And Delivery Foundation": {
            "Backend work:": [
                "Define v4.0 bounded contexts and module names so later enterprise services map cleanly to the target architecture.",
                "Create placeholder contracts for profiling, masking, ABAC, LMS, SIEM, catalog, DQ, DLP, Open Data portal, and FOI adapters.",
            ],
            "Frontend work:": [
                "Add route placeholders for governance operating model, awareness operations, DQ operations, security/compliance, and business value workspaces.",
            ],
            "Acceptance:": [
                "v4.0 scope is traceable to modules, routes, and future service boundaries without forcing microservices in Sprint 0.",
            ],
        },
        "Sprint 1: Authentication, Users, Roles, And App Shell": {
            "Backend work:": [
                "Seed the v4.0 role taxonomy: System Admin, DMO Admin, Data Owner, Business Steward, Technical Steward, Operational Data Steward, Project Data Steward, Enterprise Data Steward, DQ Steward, Privacy Steward/DPO, Security Reviewer, OD Officer, FOI Officer, Auditor, and Executive.",
                "Prepare permission namespaces for role-data mapping, certification, community, masking, ABAC, and recurring compliance.",
            ],
            "Frontend work:": [
                "Expose role-aware navigation groups for Governance Operations, Data Operations, Transparency Operations, Awareness Operations, and Security/Compliance.",
            ],
            "Acceptance:": [
                "Access behavior proves that v4.0 roles can be introduced gradually without weakening existing route guards.",
            ],
        },
        "Sprint 2: Administration Master Data Part 1": {
            "Backend work:": [
                "Add configuration foundations for policy hierarchy, charter categories, decision-right levels, and domain council metadata.",
                "Add bilingual administrative fields needed by v4.0 operating-model configuration.",
            ],
            "Frontend work:": [
                "Add nontechnical configuration views for policy hierarchy, charter metadata, and decision-right lookup values.",
            ],
            "Acceptance:": [
                "Administrators can configure governance values that later drive charters, policies, RACI, and workflow escalation.",
            ],
        },
        "Sprint 3: Administration Master Data Part 2": {
            "Backend work:": [
                "Add RACI template, DGPO sizing, domain council, and decision matrix APIs.",
                "Add stewardship role categories for Business, Technical, Operational, Project, and Enterprise stewards.",
            ],
            "Frontend work:": [
                "Build RACI template, decision matrix, DGPO sizing, and domain council configuration screens.",
            ],
            "Acceptance:": [
                "DMO Admin can model the entity operating structure required by the v4.0 governance design.",
            ],
        },
        "Sprint 4: Data Asset Governance Hub MVP": {
            "Backend work:": [
                "Prepare Asset 360 extension points for DQ, DLP, PDP, NDI, risk, lineage, glossary, training, masking, and catalog write-back signals.",
            ],
            "Frontend work:": [
                "Design Asset 360 sections so v4.0 signals appear as readable business panels, not disconnected technical tabs.",
            ],
            "Acceptance:": [
                "Asset 360 can display current MVP data while clearly reserving places for v4.0 integration signals.",
            ],
        },
        "Sprint 5: Ownership Registry And Assignment Rules MVP": {
            "Backend work:": [
                "Extend ownership logic for Operational Data Steward, Project Data Steward, Enterprise Data Steward, delegation, backup, workload, effective dates, and steward scorecard inputs.",
            ],
            "Frontend work:": [
                "Show stewardship responsibilities, workload, backup/delegation, and role-specific accountability in the registry and Asset 360.",
            ],
            "Acceptance:": [
                "The registry supports the five-role stewardship model and captures why a person is accountable.",
            ],
        },
        "Sprint 6: Workflow Engine MVP": {
            "Backend work:": [
                "Add universal case-type structure, decision recording, SLA template hooks, escalation levels, and BPMN-ready workflow metadata.",
                "Prepare case links for assignment, DQ remediation, PDP review, sharing, open data, FOI, evidence review, appeals, and access certification.",
            ],
            "Frontend work:": [
                "Add case timeline, decision panel, SLA indicator, escalation state, and clear next-action guidance.",
            ],
            "Acceptance:": [
                "Every case records owner, decision, SLA, escalation state, and audit trail consistently.",
            ],
        },
        "Sprint 7: NDI Specification Registry MVP": {
            "Backend work:": [
                "Support all 12 NDI domains: DG, MCM, DQ, DSI, PDP, DAM, DCM, RMD, BIA, DVR, OD, and FOI.",
                "Add operational lifecycle metadata, evidence requirement type, maturity level, and acceptance criteria for each specification.",
            ],
            "Frontend work:": [
                "Make the NDI registry browsable by domain, lifecycle phase, evidence type, maturity level, and responsible steward role.",
            ],
            "Acceptance:": [
                "The registry can represent the v4.0 NDI domain operating models, not only the first seven domains.",
            ],
        },
        "Sprint 8: Release 1 Hardening And UAT": {
            "Testing work:": [
                "Add UAT scenarios for RACI, decision rights, domain councils, stewardship roles, NDI domain coverage, and v4.0 navigation.",
            ],
            "Acceptance:": [
                "Release 1 proves the operating-model foundation before deeper evidence, DQ, and security modules are added.",
            ],
        },
        "Sprint 9: Evidence Repository MVP": {
            "Backend work:": [
                "Add evidence chain-of-custody metadata, expiry warnings, evidence owner, reviewer, hash verification, and compliance mapping hooks.",
            ],
            "Frontend work:": [
                "Show evidence provenance, expiry risk, reviewer decision, and linked NDI/control requirement in one evidence detail view.",
            ],
            "Acceptance:": [
                "Evidence can support NDI, PDPL, NCA ECC, masking, ABAC, and audit-pack needs later without schema rework.",
            ],
        },
        "Sprint 10: NDI Scoring And Gap Analysis": {
            "Backend work:": [
                "Add v4.0 expanded maturity dimensions: organizational awareness, roles and structures, standards and processes, and value creation.",
                "Add improvement-plan records linked to gap heat maps and 6-month assessment cycles.",
            ],
            "Frontend work:": [
                "Build maturity heat maps and improvement-plan views that explain gap priority to executives and stewards.",
            ],
            "Acceptance:": [
                "Readiness and maturity scores are explainable by evidence, role structure, process adoption, and value outcomes.",
            ],
        },
        "Sprint 11: Dashboards MVP": {
            "Backend work:": [
                "Add KPI groups for the six towers: Governance, Compliance, Data, Transparency, Value, and Awareness operations.",
            ],
            "Frontend work:": [
                "Show executive-readable cards for stewardship performance, evidence readiness, DQ trends, training/certification, transparency workload, and value outcomes.",
            ],
            "Acceptance:": [
                "Dashboards reflect the v4.0 operating model and remain role-filtered for each persona.",
            ],
        },
        "Sprint 12: Training And Awareness Module": {
            "Goal: track required training by governance role.": [
                "Goal: operationalize the v4.0 Awareness Operations Tower through training, certification, continuing education, community support, and mentorship.",
            ],
            "Backend work:": [
                "Model the 4-tier training curriculum: Tier 1 Fundamentals, Tier 2 Core Skills, Tier 3 Advanced, and Tier 4 Leadership.",
                "Build certification APIs for CDS, SDS, and MDS with prerequisites, exam results, privileges, renewal, and continuing education hours.",
                "Add community-of-practice foundations: forums, knowledge base articles, expert directory, contribution points, and mentorship tracking.",
                "Connect training and certification status to steward performance scorecards.",
            ],
            "Frontend work:": [
                "Build training pathway view by role and tier.",
                "Build certification progress view for CDS, SDS, and MDS.",
                "Build community MVP with knowledge-base browsing, expert directory, and mentorship status.",
                "Show training, certification, and CE status in steward and manager dashboards.",
            ],
            "Database work:": [
                "Add certification tracks, certification attempts, CE activities, community posts/articles, expert profiles, and mentorship pairs.",
            ],
            "Acceptance:": [
                "Users can see mandatory training, certification progress, expiry/renewal needs, and recommended community support.",
                "DMO Admin can configure training and certification requirements by stewardship role.",
                "Training and certification status affects steward readiness metrics without blocking normal governance work.",
            ],
        },
        "Sprint 13: Data Quality Operations MVP": {
            "Goal: operationalize DQ issue handling.": [
                "Goal: operationalize the v4.0 Data Quality framework with six dimensions, profiling, severity SLAs, RCA, remediation, and rule lifecycle controls.",
            ],
            "Backend work:": [
                "Implement six DQ dimensions: Accuracy, Completeness, Consistency, Timeliness, Validity, and Uniqueness.",
                "Add DQ scorecards at enterprise, domain, asset, data element, rule, and issue level.",
                "Add profiling result import for column statistics, patterns, relationships, cross-column/cross-table analysis, distributions, anomalies, and rule recommendations.",
                "Add severity tiers P1-P4 with per-stage SLA timers, breach recording, and escalation triggers.",
                "Add RCA templates for 5 Whys, Fishbone, process mapping, and lineage analysis.",
                "Add DQ rule lifecycle: definition, review, approval, deployment, monitoring, versioning, impact analysis, and retirement.",
            ],
            "Frontend work:": [
                "Build DQ scorecard views by dimension, domain, asset, and rule.",
                "Build profiling summary and suggested-rule review screens.",
                "Build RCA-guided remediation forms with severity, SLA, breach, validation, and closure panels.",
                "Build DQ rule registry and approval workflow screens.",
            ],
            "Database work:": [
                "Add DQ dimensions, rules, rule versions, score results, profiling summaries, RCA records, SLA timers, breaches, and remediation evidence links.",
            ],
            "Acceptance:": [
                "DQ issues can be created/imported, scored by dimension, assigned by ownership rules, remediated through workflow, and closed with evidence.",
                "Severity, SLA, escalation, RCA, and rule lifecycle decisions are visible and audited.",
                "DQ metrics appear on Asset 360 and dashboards.",
            ],
        },
        "Sprint 14: Classification, DLP, And Access Review MVP": {
            "Backend work:": [
                "Add role-data mapping foundations for role-to-domain access, classification limits, personal-data flags, masking requirements, and policy generation.",
                "Add masking policy model for static, dynamic, tokenization, anonymization, pseudonymization, and redaction techniques.",
                "Prepare ABAC decision logging and access certification case links.",
            ],
            "Frontend work:": [
                "Build role-data mapping designer MVP and masking policy preview inside the security governance workspace.",
                "Show access risk using business-readable language and status cues.",
            ],
            "Acceptance:": [
                "Access review can explain who has access, to which domain/classification, why, and which masking or approval rule applies.",
            ],
        },
        "Sprint 15: Catalog Integration MVP": {
            "Backend work:": [
                "Add connector contracts for catalog, lineage, DQ, DLP, PDP, NDI, risk, profiling, and training signals.",
                "Add canonical mapping fields to support external IDs, source trust, synchronization status, and write-back outcomes.",
            ],
            "Frontend work:": [
                "Show integration health and source-of-truth status in the integration admin screen and Asset 360.",
            ],
            "Acceptance:": [
                "Catalog sync remains safe for MVP while supporting v4.0 connector expansion.",
            ],
        },
        "Sprint 16: Release 2 Hardening And UAT": {
            "Testing work:": [
                "UAT scenario: certification progress and renewal warning.",
                "UAT scenario: community knowledge-base lookup and mentorship assignment.",
                "UAT scenario: DQ profiling result produces a suggested rule.",
                "UAT scenario: severity SLA breach triggers escalation.",
                "UAT scenario: role-data mapping supports access-review explanation.",
            ],
            "Acceptance:": [
                "Release 2 demonstrates operational awareness, DQ remediation, security review, and evidence readiness as one connected flow.",
            ],
        },
        "Sprint 17: Open Data Candidate Registry": {
            "Backend work:": [
                "Add candidate linkage to classification, DQ readiness, personal-data assessment, ownership, and publication-value indicators.",
            ],
            "Frontend work:": [
                "Show publication eligibility signals in language usable by ODIAO and data owners.",
            ],
            "Acceptance:": [
                "Open Data candidates are governed by ownership, classification, DQ, and ODIAO accountability.",
            ],
        },
        "Sprint 18: Open Data Assessment And Approval Workflow": {
            "Backend work:": [
                "Add anonymization, aggregation, DQ acceptability, metadata completeness, and privacy/legal review checkpoints.",
            ],
            "Frontend work:": [
                "Build a risk-led assessment wizard that explains blockers and required approvals clearly.",
            ],
            "Acceptance:": [
                "Open Data approvals produce traceable evidence for NDI and transparency compliance.",
            ],
        },
        "Sprint 19: Open Data Publication, Review, And Usage Monitoring": {
            "Backend work:": [
                "Add review cadence, usage signal, retirement/update decision, and value-measurement hooks.",
            ],
            "Acceptance:": [
                "Published data remains reviewed, measured, and governed after publication.",
            ],
        },
        "Sprint 20: FOI Request Registry And Intake": {
            "Backend work:": [
                "Prepare FOI channels for web, email, CRM, call center, and manual entry with consistent SLA start dates and identity/contact validation.",
            ],
            "Frontend work:": [
                "Make intake simple for officers while retaining enough structure for legal, privacy, and disclosure decisions.",
            ],
            "Acceptance:": [
                "FOI requests enter one governed queue regardless of intake channel.",
            ],
        },
        "Sprint 21: FOI Review, Decision, Disclosure, And Appeals": {
            "Backend work:": [
                "Add classification/exemption evidence, decision templates, partial-disclosure handling, appeal workflow, and disclosure audit trail.",
            ],
            "Acceptance:": [
                "FOI decisions are defensible, evidenced, time-bound, and appeal-ready.",
            ],
        },
        "Sprint 22: PDP Privacy Operations MVP": {
            "Backend work:": [
                "Add Privacy by Design phase gates across requirements, design, development, testing, and deployment.",
                "Add lawful-basis mapping, RoPA review, DPIA review, cross-border transfer monitoring, DSR queue, breach escalation, and consent/retention links.",
            ],
            "Frontend work:": [
                "Build Privacy by Design gate review screens and a compliance calendar view for quarterly/annual PDP obligations.",
            ],
            "Acceptance:": [
                "Privacy work is managed as operational cases with evidence, gates, reviewers, and regulatory due dates.",
            ],
        },
        "Sprint 23: Data Sharing And Integration Governance MVP": {
            "Backend work:": [
                "Link sharing decisions to role-data mapping, classification, masking, consent/legal basis, and agreement renewal.",
            ],
            "Acceptance:": [
                "Data sharing requests can show access purpose, legal basis, classification risk, and required controls.",
            ],
        },
        "Sprint 24: ODIAO Cockpit And Release 3 UAT": {
            "Testing work:": [
                "UAT scenario: ODIAO views open data and FOI backlog, SLA risk, required evidence, and unresolved privacy/security blockers.",
            ],
            "Acceptance:": [
                "The ODIAO cockpit gives one operational view of transparency workload, risks, and evidence readiness.",
            ],
        },
        "Sprint 25: Reporting And Export Framework": {
            "Backend work:": [
                "Add report models for six-tower operations, maturity heat maps, business impact metrics, user surveys, ROI attribution, DQ scorecards, training/certification, and recurring compliance.",
            ],
            "Frontend work:": [
                "Build report filters that executives and stewards can understand without knowing database fields.",
            ],
            "Acceptance:": [
                "Reports can explain governance performance, compliance posture, operational workload, and business value.",
            ],
        },
        "Sprint 26: NDI Audit Pack Generation": {
            "Backend work:": [
                "Include NCA ECC, PDPL, masking, ABAC, recurring compliance, DQ, training/certification, and stewardship evidence hooks in the audit-pack model.",
            ],
            "Frontend work:": [
                "Show audit-pack readiness by domain, control family, evidence status, owner, expiry, and unresolved blocker.",
            ],
            "Acceptance:": [
                "Audit packs are traceable enough for NDI, PDPL, NCA ECC, and internal governance review.",
            ],
        },
        "Sprint 27: Extended Domains Part 1: MCM, RMD, DAM": {
            "Backend work:": [
                "Add MDM matching candidates, match scoring, 5-step match resolution, survivorship rules, source-trust ranking, and golden-record lifecycle.",
                "Add DAM review fields for architecture decisions and lineage impact.",
            ],
            "Frontend work:": [
                "Build nontechnical match-resolution and survivorship review views.",
            ],
            "Acceptance:": [
                "MCM, RMD, and DAM flows support matching, survivorship, architecture review, and evidence generation.",
            ],
        },
        "Sprint 28: Extended Domains Part 2: DCM, BIA, DVR": {
            "Backend work:": [
                "Add business glossary governance with BDE definition, approval workflow, versioning, review/recertification, and health metrics.",
                "Add business lineage with process mapping, business-to-technical lineage bridge, visualization, and impact analysis.",
                "Add data asset valuation, use-case governance, ROI attribution, and data-user survey support.",
            ],
            "Frontend work:": [
                "Build glossary, business lineage, value dashboard, survey, and data asset valuation views.",
            ],
            "Acceptance:": [
                "DCM, BIA, and DVR flows connect governance to business meaning, lineage, and measurable value.",
            ],
        },
        "Sprint 29: Notifications, Escalations, And KSA Business Days": {
            "Backend work:": [
                "Implement four-level governance escalation: Domain Council, Data Stewardship Council, Data Governance Board, and Executive Steering Committee.",
                "Add recurring compliance calendar templates for processing activity review, cross-border transfer monitoring, annual DPIA review, and monthly DQ scorecard review.",
                "Add KSA business-day and holiday-aware SLA calculations with reminders and escalation penalties.",
            ],
            "Frontend work:": [
                "Build escalation dashboard and recurring compliance calendar with overdue, at-risk, and completed states.",
            ],
            "Acceptance:": [
                "Escalation, reminders, and compliance calendar tasks behave consistently across governance, DQ, privacy, and access workflows.",
            ],
        },
        "Sprint 30: Integration Hub Hardening": {
            "Backend work:": [
                "Formalize adapters for Catalog, DQ, DLP, ITSM, Open Data portal, FOI channels, LMS, SIEM, IAM/SSO, data profiling, masking, and ABAC decision logging.",
                "Add canonical data model mapping, webhook receiver, retry, dead-letter queue, reconciliation, and integration health metrics.",
            ],
            "Frontend work:": [
                "Build integration monitor with failed events, retries, source mapping, connector health, and reconciliation summary.",
            ],
            "Acceptance:": [
                "Failed integrations are visible, recoverable, auditable, and do not compromise core governance data.",
            ],
        },
        "Sprint 31: Security And Audit Hardening": {
            "Backend work:": [
                "Add ABAC policy preview/logging, OPA-ready decision records, Vault-ready secret patterns, mTLS deployment guidance, masking audit logs, and hash-chain audit verification.",
                "Harden sensitive evidence, masking, classification, DSR, FOI, and access-review authorization rules.",
            ],
            "Frontend work:": [
                "Build audit views for ABAC decisions, masking operations, sensitive data access, and evidence chain-of-custody.",
            ],
            "Acceptance:": [
                "Security and audit behavior aligns with v4.0 NCA ECC, PDPL, Zero Trust, masking, and ABAC requirements.",
            ],
        },
        "Sprint 32: Performance, Search, And Production Readiness": {
            "Backend work:": [
                "Add readiness for search indexing, cache, message bus, scheduled jobs, partitioning strategy, object storage, observability, backup/restore, and DR runbook alignment.",
            ],
            "Operations work:": [
                "Load-test high-volume areas: assets, DQ issues, evidence, audit logs, workflow cases, notifications, profiling results, and recurring compliance tasks.",
            ],
            "Acceptance:": [
                "The platform is ready for realistic UAT volume, production-like monitoring, and controlled deployment planning.",
            ],
        },
        "Sprint 33: Owner Recommendation Improvements": {
            "Backend work:": [
                "Add profiling-driven DQ rule recommendations and ownership recommendation explanations based on domain, asset, workload, certification, historical performance, and conflict patterns.",
            ],
            "Frontend work:": [
                "Show recommendation reasons, confidence, steward readiness, and DMO feedback without making the screen feel algorithmic or opaque.",
            ],
            "Acceptance:": [
                "Recommendations are explainable and remain overrideable with justification.",
            ],
        },
        "Sprint 34: Predictive SLA And Arabic FOI Classification Prototype": {
            "Backend work:": [
                "Prototype predictive SLA risk across workflow, DQ, FOI, privacy, and recurring compliance tasks.",
                "Prototype Arabic/English classification support for FOI and sensitive-data review with mandatory human override.",
            ],
            "Acceptance:": [
                "Predictive outputs support decisions but never replace accountable human approval.",
            ],
        },
        "Sprint 35: Final UAT, Security Review, And Documentation": {
            "Testing work:": [
                "Run v4.0 UAT regression across six towers, 12 NDI domains, training/certification, DQ profiling/RCA, role-data mapping, masking, ABAC, recurring compliance, glossary, lineage, MDM, PDP, OD, FOI, DSI, audit packs, and reports.",
            ],
            "Documentation work:": [
                "Update runbooks, support model, acceptance evidence, tool-selection notes, environment strategy, deployment checklist, and user guides.",
            ],
            "Acceptance:": [
                "All v4.0 critical capabilities are either delivered, deferred with rationale, or documented as production-phase dependencies.",
            ],
        },
        "Sprint 36: Go-Live, Stabilization, And Handover": {
            "Delivery work:": [
                "Execute production-readiness checklist for environments, backups, monitoring, DR, support tiers, incident response, access review, data retention, and hypercare.",
            ],
            "Acceptance:": [
                "The handover package includes v4.0 operating model, architecture, security, compliance, UAT, support, and roadmap traceability.",
            ],
        },
    }

    replacements_within = {
        "Sprint 12: Training And Awareness Module": {
            "Goal: track required training by governance role.": "Goal: operationalize the v4.0 Awareness Operations Tower through training, certification, continuing education, community support, and mentorship.",
        },
        "Sprint 13: Data Quality Operations MVP": {
            "Goal: operationalize DQ issue handling.": "Goal: operationalize the v4.0 Data Quality framework with six dimensions, profiling, severity SLAs, RCA, remediation, and rule lifecycle controls.",
        },
    }

    for sprint, replacements in replacements_within.items():
        start, end = find_sprint_bounds(doc, sprint)
        for old, new in replacements.items():
            idx = find_first(doc, old, start=start, end=end)
            replace_text(doc.paragraphs[idx], new)

    for sprint, marker_lines in additions.items():
        for marker, lines in marker_lines.items():
            if marker.startswith("Goal:"):
                continue
            insert_after_within_sprint(doc, sprint, marker, lines)

    insert_after_exact(
        doc,
        "Definition Of Done For Every Sprint",
        [
            "Every sprint maps delivered capabilities to the relevant v4.0 volume, operating tower, NDI domain, security/compliance requirement, and user persona.",
            "Every sprint includes a short traceability note showing which v4.0 requirements are delivered, deferred, or intentionally mocked.",
            "Every sprint preserves nontechnical usability: executive summaries, operator queues, steward tasks, and admin configuration must be readable without engineering context.",
        ],
    )

    insert_after_exact(
        doc,
        "Cross-Cutting Engineering Standards",
        [
            "Keep v4.0 enterprise integrations behind stable adapter interfaces so the MVP can use mock/CSV flows without blocking future production connectors.",
            "Treat masking, ABAC, audit, evidence chain-of-custody, and recurring compliance as cross-cutting controls; do not implement them as isolated screens only.",
            "Use event-friendly design for workflow, evidence, DQ, profiling, masking, audit, certification, and notification changes even while the first implementation remains a modular monolith.",
        ],
    )

    doc.core_properties.title = "DGOP v4.0 Sprint Details Update"
    doc.core_properties.subject = "Updated sprint roadmap aligned to DGOP v4.0 enhanced specification volumes"
    doc.core_properties.keywords = "DGOP, sprint roadmap, v4.0, data governance, NDI, SDAIA, DGA, NCA ECC, PDPL"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
