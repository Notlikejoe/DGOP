from __future__ import annotations

import hashlib
import sys
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(r"C:\Users\Youss\OneDrive\Documents\DGOP")
sys.path.insert(0, str(ROOT / "_analysis"))

import create_all_sprints_word as base  # noqa: E402


ZIP_PATH = Path(r"C:\Users\Youss\Downloads\DGOP_v5.0_ENHANCED.zip")
V5_EXTRACTED = ROOT / "_analysis" / "v5_zip_compare_20260716"
WORK_DIR = ROOT / "_analysis" / "v5_integrated_sprint_word_work"
FINAL_DOCX = ROOT / "DGOP_v5_Integrated_Sprint_Roadmap.docx"
ARTIFACT_MD = WORK_DIR / "artifact.md"


V5_VOLUME_GAPS = [
    {
        "volume": "Volume 1 - Executive Business Design",
        "missing": "Current-state challenges, quantified benefits, KPI formulas, governance committee responsibilities, DGSC/Data Council/working group charters, operating ceremonies, DMO responsibilities, six-tower operating procedures, charter management, policy lifecycle, DGPO sizing, domain council structure, community of practice governance, 4-tier training, and 3-level certification.",
        "integration": "Expanded Sprint 12, Sprint 29, Sprint 33, and new Sprint 37.",
    },
    {
        "volume": "Volume 2 - Functional Design",
        "missing": "Detailed page and data-field requirements, workflow template management, SLA template management, universal case management, case types and lifecycle, task management, notification and escalation configuration, delegation and backup, workload management, approval history, effective-date rules, and detailed integration points.",
        "integration": "Expanded Sprint 5, Sprint 6, Sprint 29, Sprint 32, and new Sprint 38.",
    },
    {
        "volume": "Volume 3 - NDI Domain Models",
        "missing": "Full per-domain operational models for DG, MCM, RMD, DAM, DCM, BIA, and DVR; NDI specification mapping; evidence requirements; evidence quality dimensional tracking; KPI and RACI matrices; BPMN workflows; metadata certification; glossary approval; MDM match resolution; golden record and survivorship concepts.",
        "integration": "Expanded Sprint 7, Sprint 10, Sprint 26, Sprint 27, Sprint 28, and new Sprint 39.",
    },
    {
        "volume": "Volume 4 - Technical Design",
        "missing": "Detailed DDD bounded-context map, anti-corruption layer, shared kernel, service decomposition, profiling engine, matching engine, ABAC policy engine, data masking service, unified search service, notification service, BPM service, service mesh/Istio, KSA cloud recommendations, deployment topology, data-flow patterns, performance architecture, caching, and resilience patterns.",
        "integration": "Expanded Sprint 15, Sprint 30, Sprint 31, Sprint 32, and new Sprint 40.",
    },
    {
        "volume": "Volume 5 - Security and Compliance",
        "missing": "Role-data mapping designer details, Open Policy Agent/OPA ABAC details, HashiCorp Vault, mTLS, Privacy by Design SDLC controls, enhanced search security, secure error handling and notification security, masking controls, evidence chain of custody, audit log integrity, SIEM monitoring, incident response, recurring compliance calendar, and ECC/PDPL/DSP control crosswalks.",
        "integration": "Expanded Sprint 14, Sprint 22, Sprint 31, Sprint 35, and new Sprint 41; dedicated enterprise error handling in Sprint 43.",
    },
    {
        "volume": "Volume 6 - Procurement and Delivery",
        "missing": "DEV/TEST/UAT/PRE-PROD/PROD/DR environment strategy, module-level acceptance criteria, procurement and delivery governance, production gates, disaster recovery expectations, support model, hypercare, and measurable performance targets.",
        "integration": "Expanded Sprint 32, Sprint 35, Sprint 36, and new Sprint 42.",
    },
]


V5_ACCEPTANCE_TARGETS = [
    "Asset search supports Arabic and English full-text queries across 1 million assets with p95 results under 500ms.",
    "Asset 360 page loads in under 2 seconds for assets with 50+ related workflow cases and 20+ evidence files.",
    "Bulk asset import processes 10,000 assets in under 10 minutes with a validation error report.",
    "Workflow engine processes 1,000 concurrent cases with task assignment latency under 1 second.",
    "Timer jobs fire within +/-30 seconds of scheduled time and use KSA business days.",
    "Workflow recovery from durable event log completes in under 5 minutes after restart.",
    "Dead-letter queue retries failed service tasks 3 times with 5-second exponential backoff before manual intervention alert.",
    "Personal-data flag accuracy is at least 99 percent against DLP scan results and catalog metadata.",
]


V5_WORKFLOW_TYPES = [
    "Owner Assignment",
    "Steward Assignment",
    "Metadata Certification",
    "DQ Remediation",
    "DLP Incident Review",
    "Data Sharing Request",
    "Open Data Candidate Review",
    "FOI Request Review",
    "FOI Appeal",
    "Personal Data Review",
    "Classification Approval",
    "Access Review",
    "NDI Evidence Review",
    "Steward Assessment",
    "Charter Approval",
    "Policy Lifecycle",
    "Glossary Term Approval",
    "MDM Match Resolution",
]


SPRINT_ENHANCEMENTS = OrderedDict(
    [
        (
            5,
            [
                "Add delegation, backup stewardship, steward workload management, approval history, and effective-date operating controls from Volume 2.",
                "Add recommendation explanation inputs for certification status, historical performance, workload, and conflict patterns from Volume 1 and Volume 2.",
            ],
        ),
        (
            6,
            [
                "Treat universal case management as a platform layer, not only a workflow screen.",
                "Add BPMN 2.0 alignment and the full 18 workflow-type backlog from Volume 6.",
            ],
        ),
        (
            7,
            [
                "Add charter management lifecycle, policy lifecycle, DGPO sizing, domain council structure, and evidence-quality dimensional tracking from Volume 3.",
            ],
        ),
        (
            12,
            [
                "Expand awareness to include 4-tier training, 3-level certification (CDS/SDS/MDS), mentorship, continuing education, and community of practice governance.",
            ],
        ),
        (
            14,
            [
                "Expand security governance to include masking service design, ABAC/OPA policy evaluation, search security, and secure decision logging.",
            ],
        ),
        (
            15,
            [
                "Expand integrations with anti-corruption layer patterns, source trust, catalog/DQ/DLP/PDP/NDI/risk/training adapter contracts, and production adapter readiness.",
            ],
        ),
        (
            26,
            [
                "Add ECC/PDPL/DSP control crosswalk output into audit packs, not only NDI evidence packaging.",
            ],
        ),
        (
            27,
            [
                "Expand MCM, RMD, and DAM with detailed NDI lifecycle, evidence, KPI, RACI, golden record, survivorship, and architecture decision requirements.",
            ],
        ),
        (
            28,
            [
                "Expand DCM, BIA, and DVR with detailed NDI lifecycle, evidence, KPI, RACI, glossary approval, metadata certification, and business value traceability.",
            ],
        ),
        (
            29,
            [
                "Add four-level governance escalation path and recurring compliance calendar rules from Volume 5 and Volume 6.",
            ],
        ),
        (
            30,
            [
                "Add unified search service, notification service, BPM service, service mesh/Istio, and integration resilience details from Volume 4.",
            ],
        ),
        (
            31,
            [
                "Add Vault, mTLS, SIEM, audit-chain integrity, incident response, secure errors, OPA/ABAC, and Privacy by Design SDLC security gates.",
            ],
        ),
        (
            32,
            [
                "Add explicit nonfunctional targets for search, Asset 360, imports, workflow concurrency, timer accuracy, event-log recovery, and dead-letter handling.",
            ],
        ),
        (
            35,
            [
                "Add module-level acceptance gates and traceability to v5 volumes, NDI domains, ECC/PDPL controls, and procurement requirements.",
            ],
        ),
        (
            36,
            [
                "Add DEV/TEST/UAT/PRE-PROD/PROD/DR environment handover, DR runbook, hypercare, support model, and production acceptance package.",
            ],
        ),
    ]
)


EXTRA_SPRINTS = [
    {
        "number": 37,
        "title": "V5 Business Operating Model Gap Closure",
        "release": "Release 6: V5 Specification Closure And Enterprise Readiness",
        "goal": "Convert the detailed v5 business operating model into sprint-ready governance operations.",
        "sections": OrderedDict(
            [
                (
                    "Scope",
                    [
                        "Document and implement DGSC, Data Council, working group, DMO, domain council, and community of practice responsibilities.",
                        "Add charter management lifecycle, policy lifecycle management, decision rights, operating ceremonies, and DGPO sizing guidance.",
                        "Map business benefits and KPI formulas into executive dashboards and release acceptance evidence.",
                    ],
                ),
                (
                    "Acceptance",
                    [
                        "Governance bodies, ceremonies, and decision flows are visible in the product or documented in AppDocs/runbooks.",
                        "Executive KPIs trace to formulas, data sources, owners, and evidence requirements.",
                        "Business operating model gaps from Volume 1 are closed or explicitly deferred.",
                    ],
                ),
            ]
        ),
    },
    {
        "number": 38,
        "title": "V5 Functional Configuration And Universal Case Management",
        "release": "Release 6: V5 Specification Closure And Enterprise Readiness",
        "goal": "Close detailed functional requirements around configurable workflow, cases, assignments, and operator screens.",
        "sections": OrderedDict(
            [
                (
                    "Scope",
                    [
                        "Add workflow template management, SLA template management, case type registry, configurable notification rules, and escalation templates.",
                        "Add delegation, backup steward assignment, steward workload management, approval history, and effective-date controls.",
                        "Document page-level field requirements, validation rules, role actions, empty/error states, and API contracts for critical modules.",
                    ],
                ),
                (
                    "Acceptance",
                    [
                        "Universal case management supports the full v5 workflow backlog and remains understandable to nontechnical operators.",
                        "Functional design tables from Volume 2 are mapped to implemented screens, API contracts, or explicit backlog items.",
                    ],
                ),
            ]
        ),
    },
    {
        "number": 39,
        "title": "V5 NDI Domain Models And Evidence Quality Closure",
        "release": "Release 6: V5 Specification Closure And Enterprise Readiness",
        "goal": "Make the detailed NDI domain operating models traceable to workflows, evidence, KPIs, and RACI.",
        "sections": OrderedDict(
            [
                (
                    "Scope",
                    [
                        "Expand DG, MCM, RMD, DAM, DCM, BIA, and DVR domain models with lifecycle, evidence, KPI, RACI, and workflow requirements.",
                        "Add evidence quality dimensional tracking, metadata certification, glossary term approval, MDM match resolution, golden record, and survivorship details.",
                        "Produce a domain-by-domain traceability matrix from Volume 3 to routes, workflows, API models, and audit-pack outputs.",
                    ],
                ),
                (
                    "Acceptance",
                    [
                        "Every Volume 3 domain section is mapped to an implemented capability, a planned sprint item, or a formal deferral.",
                        "Evidence quality and RACI expectations are visible to stewards, auditors, and executives.",
                    ],
                ),
            ]
        ),
    },
    {
        "number": 40,
        "title": "V5 Technical Architecture And Platform Services Closure",
        "release": "Release 6: V5 Specification Closure And Enterprise Readiness",
        "goal": "Close enterprise technical architecture gaps that were only summarized in the earlier roadmap.",
        "sections": OrderedDict(
            [
                (
                    "Scope",
                    [
                        "Document bounded contexts, context map, anti-corruption layer, shared kernel, service catalog, and deployment topology.",
                        "Plan profiling engine, matching engine, ABAC policy engine, data masking service, unified search service, notification service, and BPM service as platform services.",
                        "Add KSA cloud recommendation, service mesh/Istio option, resilience patterns, caching strategy, search indexing strategy, and integration data-flow patterns.",
                    ],
                ),
                (
                    "Acceptance",
                    [
                        "Architecture decisions are recorded with tradeoffs, risks, and staged implementation plan.",
                        "Each Volume 4 technical service has an owner sprint, API/data contract, and operational readiness path.",
                    ],
                ),
            ]
        ),
    },
    {
        "number": 41,
        "title": "V5 Security, Compliance, And Control Crosswalk Closure",
        "release": "Release 6: V5 Specification Closure And Enterprise Readiness",
        "goal": "Make security, privacy, and compliance controls explicit, testable, and auditable.",
        "sections": OrderedDict(
            [
                (
                    "Scope",
                    [
                        "Add detailed OPA/ABAC, role-data mapping designer, HashiCorp Vault, mTLS, SIEM monitoring, incident response, secure error handling, and enhanced search security requirements.",
                        "Add Privacy by Design SDLC gates, masking controls, evidence chain-of-custody, audit log integrity checks, and recurring compliance calendar responsibilities.",
                        "Create ECC-to-DGOP, PDPL 2023, ECC-to-DSP, and security checklist traceability outputs.",
                    ],
                ),
                (
                    "Acceptance",
                    [
                        "Every critical security/compliance control from Volume 5 is mapped to an implementation, test, audit artifact, or accepted deferral.",
                        "Security controls are enforced server-side and represented in release readiness evidence.",
                    ],
                ),
            ]
        ),
    },
    {
        "number": 42,
        "title": "V5 Procurement, Environments, Performance, And Production Acceptance",
        "release": "Release 6: V5 Specification Closure And Enterprise Readiness",
        "goal": "Convert procurement, environment, and measurable production acceptance requirements into a final readiness package.",
        "sections": OrderedDict(
            [
                (
                    "Scope",
                    [
                        "Define DEV, TEST, UAT, PRE-PROD, PROD, and DR environment strategy with entry/exit criteria, data controls, support responsibilities, and DR expectations.",
                        "Add module-level acceptance criteria for administration, ownership, Asset 360, workflow, NDI, training, DQ/security, integrations, reports, and audit packs.",
                        "Run or plan acceptance targets: 1M asset search under 500ms p95, Asset 360 under 2 seconds, 10,000 asset import under 10 minutes, 1,000 concurrent workflow cases, timer accuracy within +/-30 seconds, recovery under 5 minutes, and DLQ retries.",
                    ],
                ),
                (
                    "Acceptance",
                    [
                        "Procurement and delivery expectations from Volume 6 are traceable to go-live gates.",
                        "Production readiness package includes environment, performance, DR, support, hypercare, known issues, and sign-off evidence.",
                    ],
                ),
            ]
        ),
    },
    {
        "number": 43,
        "title": "Enterprise Error Experience And Observability Hardening",
        "release": "Release 7: Enterprise Stabilization And Demo Trust",
        "goal": "Turn DGOP error handling into a consistent, nontechnical, secure, and auditable platform experience.",
        "sections": OrderedDict(
            [
                (
                    "Scope",
                    [
                        "Standardize backend error envelopes with stable error codes, request/correlation IDs, safe messages, route context, and consistent handling for validation, authorization, conflict, rate-limit, service-unavailable, import, and unexpected failures.",
                        "Add a shared Angular error interpretation layer so forms, tables, dashboards, workflow actions, file imports, login/session expiry, and API-down states show clear bilingual guidance instead of raw or generic failures.",
                        "Create reusable page-level and field-level error states with retry actions, degraded-mode messaging, and nontechnical next steps for governance operators, stewards, auditors, and administrators.",
                        "Add structured error telemetry, safe redaction rules, audit-log linkage for governance-impacting failures, and operational reports for recurring errors without exposing sensitive data.",
                    ],
                ),
                (
                    "Acceptance",
                    [
                        "All critical API failures return a predictable safe error body with status, stable code, user-safe message, timestamp, path, and request/correlation ID.",
                        "The frontend maps 400, 401, 403, 404, 409, 422, 429, 500, network, timeout, upload/import, and stale-session failures to clear English and Arabic UX states.",
                        "Validation errors appear beside the affected field where possible; row-level CSV/import failures are downloadable and reviewable without losing successful records.",
                        "Permission and session errors explain what happened and how to recover without exposing roles, tokens, stack traces, SQL details, or internal service names.",
                        "Workflow, evidence, DQ, security, integrations, access control, and dashboard pages have tested empty, loading, degraded, retry, and failure states in light mode, dark mode, English, and Arabic RTL.",
                    ],
                ),
                (
                    "Risks And Mitigations",
                    [
                        "Risk: over-generic messages slow support. Mitigation: separate user-safe text from operator diagnostics through request IDs and internal logs.",
                        "Risk: raw backend messages leak sensitive implementation details. Mitigation: allowlist public error codes/messages and redact logs before telemetry export.",
                        "Risk: toast noise overwhelms operators. Mitigation: use page-level errors for blocked work, inline errors for form fields, and toasts only for short-lived action results.",
                        "Risk: retrofitting every module misses edge cases. Mitigation: add an error-state QA matrix covering auth, governance, admin, workflow, imports, evidence, and integrations.",
                    ],
                ),
            ]
        ),
    },
]


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def add_cover(doc: Document, sprints: list[dict]) -> None:
    for _ in range(6):
        base.add_para(doc, "")
    title1 = doc.add_paragraph(style="Title")
    title1.add_run("DGOP v5").font.color.rgb = RGBColor.from_string(base.MID_BLUE)
    title2 = doc.add_paragraph(style="Title")
    run = title2.add_run("Integrated Sprint Roadmap")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(base.BLUE)
    base.add_para(doc, "Updated sprint roadmap including v5/v4.1 enhanced specification gaps from all six source volumes.", "normal")
    base.add_para(doc, "")

    meta = doc.add_table(rows=1, cols=3)
    values = [
        ("STATUS", "Updated / Gap Integrated"),
        ("OWNER", "DGOP Team"),
        ("LAST UPDATED", "July 16, 2026"),
    ]
    for cell, (label, value) in zip(meta.rows[0].cells, values):
        p = cell.paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(7)
        r.font.color.rgb = RGBColor.from_string(base.BLUE)
        p.add_run(f"\n{value}")
    base.format_table(meta, header=False)
    base.add_para(doc, "")
    info = doc.add_table(rows=4, cols=2)
    rows = [
        ("Authors", "Codex with DGOP v4 roadmap and v5 enhanced source package"),
        ("Reviewers", "DGOP senior stakeholders"),
        ("Related docs", "DGOP_v5.0_ENHANCED.zip; DGOP_All_Sprints_System_Design_Roadmap.docx"),
        ("Scope", f"Sprints {sprints[0]['number']} through {sprints[-1]['number']} with explicit v5 gap closure."),
    ]
    for row, (k, v) in zip(info.rows, rows):
        row.cells[0].text = k
        row.cells[1].text = v
    base.format_table(info, header=False)
    doc.add_page_break()


def add_intro(doc: Document, sprints: list[dict]) -> None:
    doc.add_heading("1. Abstract", level=1)
    base.add_para(
        doc,
        "This document updates the DGOP sprint roadmap by integrating the missing details identified from the enhanced source package. It preserves the original sprint sequence, adds explicit v5 traceability, and creates dedicated gap-closure sprints for areas that were previously summarized too lightly.",
    )
    base.add_para(
        doc,
        "The result is a roadmap that can be used by senior data governance, security, engineering, procurement, and delivery stakeholders to understand what must be built, verified, documented, or formally deferred before production readiness.",
    )

    doc.add_heading("2. Goals and Non-Goals", level=1)
    table = doc.add_table(rows=5, cols=2)
    rows = [
        ("Goals", "Non-goals"),
        ("Integrate missing v5 content into a sprint-level delivery plan.", "This document does not reproduce all six source volumes word-for-word."),
        ("Show where every major missing area belongs in the roadmap.", "This document does not certify that the current app already implements every future item."),
        ("Add explicit gap-closure sprints for architecture, security, NDI, functional, operating-model, and delivery gaps.", "This document does not replace QA test cases, procurement documents, or final sign-off evidence."),
        ("Keep the System Design template style and produce a professional Word artifact.", "This document does not expose secrets or local credentials."),
    ]
    for row, values in zip(table.rows, rows):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]
    base.format_table(table)

    doc.add_heading("3. Source Package Coverage Matrix", level=1)
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ["Source volume", "Previously missing or too light", "Integrated into roadmap"]):
        cell.text = text
    for item in V5_VOLUME_GAPS:
        row = table.add_row().cells
        row[0].text = item["volume"]
        row[1].text = item["missing"]
        row[2].text = item["integration"]
    base.format_table(table)

    doc.add_heading("4. Explicit V5 Acceptance Targets", level=1)
    base.add_bullets(doc, V5_ACCEPTANCE_TARGETS)

    doc.add_heading("5. Required Workflow Coverage", level=1)
    base.add_para(doc, "The workflow engine roadmap must account for all 18 governance workflow types from the enhanced delivery criteria:")
    base.add_bullets(doc, V5_WORKFLOW_TYPES)


def add_enhancement_matrix(doc: Document) -> None:
    doc.add_heading("6. Existing Sprint Enhancement Matrix", level=1)
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ["Sprint", "Roadmap area", "V5 additions now required"]):
        cell.text = text
    area_by_sprint = {
        5: "Ownership and stewardship",
        6: "Workflow engine",
        7: "NDI foundation",
        12: "Training and awareness",
        14: "Security governance",
        15: "Catalog and integrations",
        26: "Audit packs",
        27: "Extended domains part 1",
        28: "Extended domains part 2",
        29: "Notifications and escalations",
        30: "Integration hub",
        31: "Security hardening",
        32: "Performance and readiness",
        35: "Final UAT",
        36: "Go-live and handover",
    }
    for sprint_no, additions in SPRINT_ENHANCEMENTS.items():
        row = table.add_row().cells
        row[0].text = f"Sprint {sprint_no}"
        row[1].text = area_by_sprint[sprint_no]
        row[2].text = " ".join(additions)
    base.format_table(table)


def add_updated_sprint_catalog(doc: Document, sprints: list[dict]) -> None:
    doc.add_heading("7. Updated Sprint Catalog Overview", level=1)
    table = doc.add_table(rows=1, cols=4)
    for cell, text in zip(table.rows[0].cells, ["Sprint", "Release", "Title", "Goal"]):
        cell.text = text
    for sprint in sprints:
        row = table.add_row().cells
        row[0].text = str(sprint["number"])
        row[1].text = sprint["release"].split(":", 1)[0]
        row[2].text = sprint["title"]
        row[3].text = sprint["goal"] or "See detailed section."
    base.format_table(table)


def add_sprint_details(doc: Document, sprints: list[dict]) -> None:
    doc.add_page_break()
    doc.add_heading("8. Detailed Updated Sprint Catalog", level=1)
    current_release = None
    for sprint in sprints:
        if sprint["release"] != current_release:
            current_release = sprint["release"]
            doc.add_heading(current_release, level=2)
        doc.add_heading(f"Sprint {sprint['number']}: {sprint['title']}", level=3)
        if sprint["goal"]:
            base.add_lead(doc, "Goal", sprint["goal"])
        if sprint["number"] in SPRINT_ENHANCEMENTS:
            p = doc.add_paragraph(style="normal")
            r = p.add_run("V5 additions")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(base.MID_BLUE)
            base.add_bullets(doc, SPRINT_ENHANCEMENTS[sprint["number"]])
        for label, items in sprint["sections"].items():
            if label == "Overview" and not items:
                continue
            if label == "Goal":
                continue
            p = doc.add_paragraph(style="normal")
            r = p.add_run(label if label != "Overview" else "Additional notes")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(base.MID_BLUE)
            base.add_bullets(doc, items)


def add_appendices(doc: Document, appendices: OrderedDict[str, list[str]]) -> None:
    doc.add_page_break()
    doc.add_heading("9. Roadmap-Wide Standards and Appendices", level=1)
    for heading, items in appendices.items():
        doc.add_heading(heading, level=2)
        base.add_bullets(doc, items)
    doc.add_heading("V5 Source Package Traceability Notes", level=2)
    base.add_bullets(
        doc,
        [
            "The zip file was treated as six enhanced Word specification volumes and compared against the previous sprint roadmap.",
            "Items that were detailed in the zip but only summarized in the sprint roadmap are now captured in the source package coverage matrix, existing sprint enhancements, or Sprints 37-42.",
            "Detailed source-volume tables are not copied verbatim; they should be converted into QA test cases, backlog tickets, acceptance criteria, and implementation tasks during sprint planning.",
        ],
    )


def write_artifact(sprints: list[dict], appendices: OrderedDict[str, list[str]]) -> None:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    ARTIFACT_MD.write_text(
        "\n".join(
            [
                "# DGOP v5 Integrated Sprint Roadmap Artifact",
                "",
                f"Reference: {base.REFERENCE}",
                f"Reference SHA-256: {sha256(base.REFERENCE)}",
                "Reference page system: US Letter portrait; 0.70in left/right/top, 0.62in bottom; different first page.",
                "Reference visual system: retained System Design blue hierarchy and table treatment.",
                f"Previous sprint source: {base.SOURCE}",
                f"V5 source zip: {ZIP_PATH}",
                f"V5 extracted folder: {V5_EXTRACTED}",
                f"Generated output: {FINAL_DOCX}",
                f"Sprint coverage: {sprints[0]['number']} through {sprints[-1]['number']} ({len(sprints)} sprints).",
                "Added sections: source package coverage matrix, v5 acceptance targets, required workflow coverage, existing sprint enhancement matrix, detailed updated sprint catalog, appendices.",
                f"Appendix sections retained: {', '.join(appendices.keys())}",
                "Render note: PNG render QA requires LibreOffice/soffice; if unavailable, structural DOCX audits are used.",
                "",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    sprints, appendices = base.parse_source()
    sprints = sprints + EXTRA_SPRINTS
    write_artifact(sprints, appendices)
    doc = Document(base.REFERENCE)
    base.clear_body(doc)
    base.style_document(doc)
    add_cover(doc, sprints)
    add_intro(doc, sprints)
    add_enhancement_matrix(doc)
    add_updated_sprint_catalog(doc, sprints)
    add_sprint_details(doc, sprints)
    add_appendices(doc, appendices)
    doc.core_properties.title = "DGOP v5 Integrated Sprint Roadmap"
    doc.core_properties.subject = "Updated DGOP sprint roadmap with v5 enhanced source package gaps"
    doc.core_properties.author = "Codex"
    doc.core_properties.last_modified_by = "Codex"
    doc.core_properties.keywords = "DGOP, v5, sprint roadmap, system design, data governance"
    doc.core_properties.comments = "Created from the retained System Design template reference."
    doc.save(FINAL_DOCX)
    print(FINAL_DOCX)
    print(f"sprints={len(sprints)} appendices={len(appendices)}")


if __name__ == "__main__":
    main()
